"""
Document Loader for Resilience2Relief AI
Supports multiple document formats: PDF, DOCX, TXT, CSV
Extracts metadata and preprocesses text for RAG pipeline
"""

import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import pandas as pd

# Document processing libraries
try:
    from unstructured.partition.auto import partition
    from unstructured.documents.elements import Title, NarrativeText, ListItem, Table
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False

try:
    import PyMuPDF as fitz  # PyMuPDF for PDF processing
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentMetadata:
    """Container for document metadata"""
    def __init__(self, filename: str, file_path: str, file_size: int, 
                 disaster_type: Optional[str] = None, region: Optional[str] = None,
                 document_type: Optional[str] = None, source: Optional[str] = None,
                 date_created: Optional[datetime] = None):
        self.filename = filename
        self.file_path = file_path
        self.file_size = file_size
        self.disaster_type = disaster_type
        self.region = region
        self.document_type = document_type
        self.source = source
        self.date_created = date_created or datetime.now()
        
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        return {
            "filename": self.filename,
            "file_path": self.file_path,
            "file_size": self.file_size,
            "disaster_type": self.disaster_type,
            "region": self.region,
            "document_type": self.document_type,
            "source": self.source,
            "date_created": self.date_created.isoformat() if self.date_created else None
        }

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.csv', '.md']
        
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using PyMuPDF"""
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF not available. Install with: pip install PyMuPDF")
            
        try:
            doc = fitz.open(file_path)
            text = ""
            metadata = {
                "pages": len(doc),
                "title": doc.metadata.get('title', ''),
                "author": doc.metadata.get('author', ''),
                "subject": doc.metadata.get('subject', ''),
                "creator": doc.metadata.get('creator', ''),
                "producer": doc.metadata.get('producer', ''),
                "creation_date": doc.metadata.get('creationDate', ''),
                "modification_date": doc.metadata.get('modDate', '')
            }
            
            for page in doc:
                text += page.get_text()
                
            doc.close()
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX using python-docx"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
            
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            # Basic metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or '',
                "author": core_props.author or '',
                "subject": core_props.subject or '',
                "created": core_props.created.isoformat() if core_props.created else '',
                "modified": core_props.modified.isoformat() if core_props.modified else '',
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            }
            
            return "\n\n".join(text_parts), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "encoding": "utf-8",
                "lines_count": len(text.split('\n')),
                "words_count": len(text.split()),
                "characters_count": len(text)
            }
            
            return text.strip(), metadata
            
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            metadata = {
                "encoding": "latin-1",
                "lines_count": len(text.split('\n')),
                "words_count": len(text.split()),
                "characters_count": len(text)
            }
            
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise
    
    def extract_text_from_csv(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text
            text_parts = []
            text_parts.append(f"CSV Data with {len(df)} rows and {len(df.columns)} columns")
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            
            # Add sample data
            for idx, row in df.head(10).iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            if len(df) > 10:
                text_parts.append(f"... and {len(df) - 10} more rows")
            
            metadata = {
                "rows_count": len(df),
                "columns_count": len(df.columns),
                "columns": df.columns.tolist(),
                "data_types": df.dtypes.to_dict()
            }
            
            return "\n".join(text_parts), metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_path}: {str(e)}")
            raise
    
    def extract_metadata_from_filename(self, filename: str) -> Dict[str, Any]:
        """Extract metadata from filename patterns"""
        filename_lower = filename.lower()
        
        # Disaster type detection
        disaster_types = {
            'cyclone': ['cyclone', 'hurricane', 'typhoon'],
            'earthquake': ['earthquake', 'seismic'],
            'tsunami': ['tsunami'],
            'flood': ['flood', 'flooding'],
            'drought': ['drought'],
            'volcanic': ['volcanic', 'volcano'],
            'wildfire': ['wildfire', 'fire'],
            'landslide': ['landslide']
        }
        
        disaster_type = None
        for disaster, keywords in disaster_types.items():
            if any(keyword in filename_lower for keyword in keywords):
                disaster_type = disaster
                break
        
        # Region detection (Pacific focus)
        regions = {
            'vanuatu': ['vanuatu'],
            'samoa': ['samoa'],
            'fiji': ['fiji'],
            'tonga': ['tonga'],
            'solomon_islands': ['solomon', 'solomons'],
            'papua_new_guinea': ['papua', 'png'],
            'marshall_islands': ['marshall'],
            'palau': ['palau'],
            'micronesia': ['micronesia'],
            'nauru': ['nauru'],
            'kiribati': ['kiribati'],
            'tuvalu': ['tuvalu']
        }
        
        region = None
        for reg, keywords in regions.items():
            if any(keyword in filename_lower for keyword in keywords):
                region = reg
                break
        
        # Document type detection
        doc_types = {
            'assessment': ['assessment', 'evaluation', 'analysis'],
            'recovery_plan': ['recovery', 'reconstruction', 'rebuild'],
            'funding_guide': ['funding', 'finance', 'budget', 'donor'],
            'best_practices': ['best_practices', 'guidelines', 'manual'],
            'case_study': ['case_study', 'example', 'success']
        }
        
        document_type = None
        for doc_type, keywords in doc_types.items():
            if any(keyword in filename_lower for keyword in keywords):
                document_type = doc_type
                break
        
        return {
            'disaster_type': disaster_type,
            'region': region,
            'document_type': document_type
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a single document and return text with metadata"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Extract basic file information
        file_stats = file_path.stat()
        filename_metadata = self.extract_metadata_from_filename(file_path.name)
        
        # Extract text based on file type
        try:
            if file_path.suffix.lower() == '.pdf':
                text, doc_metadata = self.extract_text_from_pdf(str(file_path))
            elif file_path.suffix.lower() == '.docx':
                text, doc_metadata = self.extract_text_from_docx(str(file_path))
            elif file_path.suffix.lower() in ['.txt', '.md']:
                text, doc_metadata = self.extract_text_from_txt(str(file_path))
            elif file_path.suffix.lower() == '.csv':
                text, doc_metadata = self.extract_text_from_csv(str(file_path))
            else:
                raise ValueError(f"Unsupported format: {file_path.suffix}")
            
            # Combine all metadata
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_path=str(file_path),
                file_size=file_stats.st_size,
                disaster_type=filename_metadata.get('disaster_type'),
                region=filename_metadata.get('region'),
                document_type=filename_metadata.get('document_type'),
                date_created=datetime.fromtimestamp(file_stats.st_ctime)
            )
            
            return {
                'text': text,
                'metadata': metadata.to_dict(),
                'document_metadata': doc_metadata,
                'word_count': len(text.split()),
                'character_count': len(text)
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Process all supported documents in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                try:
                    logger.info(f"Processing: {file_path.name}")
                    doc_data = self.process_document(str(file_path))
                    documents.append(doc_data)
                except Exception as e:
                    logger.error(f"Failed to process {file_path.name}: {str(e)}")
                    continue
        
        logger.info(f"Successfully processed {len(documents)} documents")
        return documents

# Utility functions
def clean_text(text: str) -> str:
    """Clean and preprocess extracted text"""
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s\.,;:!?()-]', ' ', text)
    
    # Remove excessive punctuation
    text = re.sub(r'[.]{3,}', '...', text)
    text = re.sub(r'[-]{2,}', '--', text)
    
    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into chunks with overlap for better context preservation"""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end < len(text):
            # Find the last sentence boundary
            last_period = text.rfind('.', start, end)
            last_newline = text.rfind('\n', start, end)
            boundary = max(last_period, last_newline)
            
            if boundary > start:
                end = boundary + 1
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        
        if start >= len(text):
            break
    
    return chunks

# Main interface
def load_documents(directory_path: str) -> List[Dict[str, Any]]:
    """Main function to load and process documents"""
    processor = DocumentProcessor()
    return processor.process_directory(directory_path)

def load_document(file_path: str) -> Dict[str, Any]:
    """Load and process a single document"""
    processor = DocumentProcessor()
    return processor.process_document(file_path)

if __name__ == "__main__":
    # Test with sample data
    import sys
    
    if len(sys.argv) > 1:
        path = sys.argv[1]
        
        if os.path.isfile(path):
            result = load_document(path)
            print(f"Processed: {result['metadata']['filename']}")
            print(f"Word count: {result['word_count']}")
            print(f"Disaster type: {result['metadata']['disaster_type']}")
            print(f"Region: {result['metadata']['region']}")
        elif os.path.isdir(path):
            results = load_documents(path)
            print(f"Processed {len(results)} documents")
            for result in results:
                print(f"- {result['metadata']['filename']} ({result['word_count']} words)")
        else:
            print(f"Path not found: {path}")
    else:
        print("Usage: python document_loader.py <file_or_directory_path>")